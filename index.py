import pdfplumber
import json
import os
from openai import OpenAI
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# --- CONFIGURATION ---
# OPENAI_API_KEY = "fd"
TEMPLATE_PATH = "ravan.docx"
OUTPUT_PATH = "Final_Generated_Resume.docx"
PDF_PATH = "test.pdf"

client = OpenAI(api_key=OPENAI_API_KEY)

def extract_text_from_pdf(pdf_path):
    """Reads raw text from the candidate's PDF"""
    print(f"Reading PDF: {pdf_path}...")
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text

def get_ai_data(resume_text):
    """Sends text to OpenAI and asks for structured JSON with smart filtering"""
    print("Sending data to AI for extraction...")
   
    system_prompt = """
    You are a professional resume parser following Zensar resume guidelines.
   
    CRITICAL RULES - SMART FILTERING:
    1. Summary MUST be 5+ lines highlighting total experience, projects, expertise
    2. Skills MUST be an ARRAY with 5-7 most relevant categories (NOT ALL)
       - Intelligently group similar skills (e.g., React+Angular+Vue → "Frontend Frameworks")
       - Each skill object has: category (string), primary (string), secondary (string)
       - Category names are DYNAMIC based on candidate's profile
    3. Experience summary: Limit to 3-5 most relevant/recent roles
    4. Certifications: Only professional/industry-recognized (max 5-7)
    5. Awards: Only significant achievements (max 3-5)
    6. Expertise areas: 2-3 key areas only
    7. Education: Only engineering degree and above with passing year
    8. Avoid generic phrases like "collaborated with team"
    9. Quantify all achievements with specific numbers
    """

    user_prompt = """
    Resume Content:
    """ + resume_text + """

    Extract and structure the data in this EXACT JSON format:
    {
      "personal_info": {
        "name": "Full Name",
        "email": "email@example.com",
        "phone": "+1234567890",
        "location": "City, Country"
      },
      "summary": "Professional summary here...",
      "education": [{"degree": "B.Tech", "year": "2018"}],
      "skills": [
        {
          "category": "Frontend Frameworks",
          "primary": "React, Angular",
          "secondary": "Vue.js"
        }
      ],
      "expertise_areas": ["Area 1", "Area 2"],
      "certifications": ["Cert 1"],
      "awards": ["Award 1"],
      "experience_summary": [
        {
          "role": "Job Title",
          "years": "2020-Present",
          "skills": "Skill 1, Skill 2",
          "achievements": ["Achievement 1", "Achievement 2"]
        }
      ]
    }
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        response_format={"type": "json_object"}
    )
   
    return json.loads(response.choices[0].message.content)


# ========== FIXED BORDER LOGIC START ==========

def add_table_borders(table):
    """
    Add borders to the ENTIRE table (tblPr).
    This forces 'insideV' (vertical columns) and 'insideH' (horizontal rows) to appear.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Create or get tblBorders element
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    # Define all 6 border types needed for a full grid
    # 'top', 'left', 'bottom', 'right' = Outer frame
    # 'insideH' = Horizontal lines between rows
    # 'insideV' = Vertical lines between columns
    border_map = {
        "top": "w:top",
        "left": "w:left",
        "bottom": "w:bottom",
        "right": "w:right",
        "insideH": "w:insideH",
        "insideV": "w:insideV"
    }

    for border_name, xml_tag in border_map.items():
        # Remove existing tag if present to avoid duplication
        existing = tblBorders.find(qn(xml_tag))
        if existing is not None:
            tblBorders.remove(existing)
        
        # Create new border element
        border = OxmlElement(xml_tag)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')       # 4 = 1/2 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # Black
        
        tblBorders.append(border)

# ========== FIXED BORDER LOGIC END ==========


def find_table_by_header(doc, header_text):
    """Find a table by searching for header text in the first row"""
    for table in doc.tables:
        first_row_text = ' '.join([cell.text for cell in table.rows[0].cells]).lower()
        if header_text.lower() in first_row_text:
            return table
    return None


def populate_skills_table(doc, skills_data):
    """Find and populate the skills table programmatically"""
    print("Populating skills table programmatically...")
    
    skills_table = find_table_by_header(doc, "category")
    if not skills_table:
        skills_table = find_table_by_header(doc, "primary")
    
    if not skills_table:
        print("Warning: Skills table not found! Skipping.")
        return
    
    num_columns = len(skills_table.rows[0].cells)
    
    # Delete existing data rows (keep only header row)
    for i in range(len(skills_table.rows) - 1, 0, -1):
        skills_table._element.remove(skills_table.rows[i]._element)
    
    # Add new rows for each skill
    for i, skill in enumerate(skills_data, 1):
        row = skills_table.add_row()
        
        if num_columns == 4:
            # 4-column: [No. | Category | Primary | Secondary]
            row.cells[0].text = str(i)
            row.cells[1].text = skill['category']
            row.cells[2].text = skill['primary']
            row.cells[3].text = skill['secondary']
        elif num_columns == 3:
            row.cells[0].text = skill['category']
            row.cells[1].text = skill['primary']
            row.cells[2].text = skill['secondary']
    
    # Apply the borders to the TABLE itself
    print("Adding borders to skills table...")
    add_table_borders(skills_table)
    
    print(f"✓ Added {len(skills_data)} rows to skills table with full grid borders")


def generate_doc(data, template_file, output_file):
    """Injects the JSON data into the Word Template"""
    print("Generating Word Document...")
   
    # Limits
    if 'skills' in data and len(data['skills']) > 7:
        data['skills'] = data['skills'][:7]
    if 'experience_summary' in data and len(data['experience_summary']) > 5:
        data['experience_summary'] = data['experience_summary'][:5]
   
    # 1. Render Template
    doc = DocxTemplate(template_file)
    doc.render(data)
    doc.save(output_file)
    
    # 2. Post-process (Table)
    print("\n" + "="*80)
    print("POST-PROCESSING: Adding skills table programmatically...")
    print("="*80)
    
    rendered_doc = Document(output_file)
    populate_skills_table(rendered_doc, data['skills'])
    rendered_doc.save(output_file)
    
    print(f"\n✓ Success! Resume saved to: {output_file}")


# --- MAIN EXECUTION ---
if __name__ == "__main__":
    if not os.path.exists(TEMPLATE_PATH):
        print(f"Error: Template file '{TEMPLATE_PATH}' not found.")
    elif not os.path.exists(PDF_PATH):
        print(f"Error: PDF file '{PDF_PATH}' not found.")
    else:
        raw_text = extract_text_from_pdf(PDF_PATH)
        resume_data = get_ai_data(raw_text)
        
        print("\n" + "="*80)
        print("EXTRACTED DATA:")
        print("="*80)
        print(json.dumps(resume_data, indent=2))
        print("="*80 + "\n")
        
        generate_doc(resume_data, TEMPLATE_PATH, OUTPUT_PATH)